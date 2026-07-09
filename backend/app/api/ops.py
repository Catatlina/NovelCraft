"""项目运营 API — 搜索/导出/版本管理 (Phase 9)"""
from __future__ import annotations

import io
import os
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import select, func, or_
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user, get_user_project
from app.db.database import get_db
from app.db.models import NovelChapter, NovelProject, User
from app.schemas import ExportRequest, ExportResponse

router = APIRouter(prefix="/api/v1/ops", tags=["ops"])


@router.get("/search")
async def search_projects(q: str = Query(default=""), status: str | None = None,
                          genre: str | None = None, db: AsyncSession = Depends(get_db),
                          user: User = Depends(get_current_user)):
    stmt = select(NovelProject).where(NovelProject.user_id == user.id)
    if q: stmt = stmt.where(or_(NovelProject.title.ilike(f"%{q}%"), NovelProject.genre.ilike(f"%{q}%")))
    if status: stmt = stmt.where(NovelProject.status == status)
    if genre: stmt = stmt.where(NovelProject.genre == genre)
    stmt = stmt.order_by(NovelProject.updated_at.desc()).limit(50)
    r = await db.execute(stmt)
    return [{"id": str(p.id), "title": p.title, "genre": p.genre, "platform": p.platform,
             "status": p.status, "total_chapters": p.total_chapters,
             "total_words": p.total_words, "updated_at": str(p.updated_at)} for p in r.scalars().all()]


# ---------------------------------------------------------------------------
# Multi-format Export
# ---------------------------------------------------------------------------

_EXPORT_DIR: Path | None = None

def _get_export_dir() -> Path:
    global _EXPORT_DIR
    if _EXPORT_DIR is None:
        _EXPORT_DIR = Path(__file__).resolve().parent.parent.parent / "exports"
        _EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return _EXPORT_DIR


def _build_chapter_lines(chapters: list[NovelChapter], project_title: str) -> list[str]:
    """Build common text lines for TXT export."""
    lines = [f"《{project_title}》\n"]
    for ch in chapters:
        lines.append(f"{'=' * 40}\n第{ch.chapter_num}章 {ch.title or ''}\n{'=' * 40}\n")
        if ch.content:
            lines.append(ch.content + "\n")
    return lines


def _export_txt(chapters: list[NovelChapter], project_title: str, encoding: str = "utf-8") -> bytes:
    """Export as plain text."""
    lines = _build_chapter_lines(chapters, project_title)
    content = "".join(lines)
    return content.encode(encoding, errors="replace")


def _export_epub(chapters: list[NovelChapter], project_title: str, include_cover: bool = False,
                 cover_text: str | None = None) -> bytes:
    """Export as EPUB using ebooklib."""
    try:
        from ebooklib import epub
    except ImportError:
        raise HTTPException(500, "ebooklib not installed. Install with: pip install ebooklib")

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(project_title)
    book.set_language("zh")
    book.add_author("NovelCraft Author")

    # Build spine chapters
    spine = ["nav"]
    epub_chapters = []

    for i, ch in enumerate(chapters):
        c = epub.EpubHtml(
            title=ch.title or f"Chapter {ch.chapter_num}",
            file_name=f"chap_{ch.chapter_num:04d}.xhtml",
            lang="zh",
        )
        content_html = f"<h1>第{ch.chapter_num}章 {ch.title or ''}</h1>"
        if ch.content:
            for para in ch.content.split("\n"):
                if para.strip():
                    content_html += f"<p>{para.strip()}</p>"
        c.content = content_html
        book.add_item(c)
        spine.append(c)
        epub_chapters.append(c)

    book.toc = epub_chapters

    # Add navigation
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # Style
    style = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content="body { font-family: serif; line-height: 1.8; } h1 { font-size: 1.5em; }",
    )
    book.add_item(style)

    book.spine = spine

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


def _export_docx(chapters: list[NovelChapter], project_title: str) -> bytes:
    """Export as DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Install with: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading(project_title, 0)
    title.alignment = 1  # centered

    for ch in chapters:
        doc.add_heading(f"第{ch.chapter_num}章 {ch.title or ''}", level=1)
        if ch.content:
            for para_text in ch.content.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    style = p.style
                    style.font.size = Pt(12)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_pdf(chapters: list[NovelChapter], project_title: str) -> bytes:
    """Export as PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise HTTPException(500, "reportlab not installed. Install with: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontSize=24, spaceAfter=12*mm)
    heading_style = ParagraphStyle("HeadingCN", parent=styles["Heading1"], fontSize=16, spaceBefore=8*mm, spaceAfter=4*mm)
    body_style = ParagraphStyle("BodyCN", parent=styles["Normal"], fontSize=11, leading=18)

    story = [Paragraph(project_title, title_style), Spacer(1, 10*mm)]

    for ch in chapters:
        story.append(Paragraph(f"第{ch.chapter_num}章 {ch.title or ''}", heading_style))
        if ch.content:
            for para_text in ch.content.split("\n"):
                if para_text.strip():
                    story.append(Paragraph(para_text.strip(), body_style))
        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


EXPORTERS = {
    "txt": _export_txt,
    "epub": _export_epub,
    "docx": _export_docx,
    "pdf": _export_pdf,
}

CONTENT_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "epub": "application/epub+zip",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

FILE_EXTENSIONS = {
    "txt": ".txt",
    "epub": ".epub",
    "docx": ".docx",
    "pdf": ".pdf",
}


@router.post("/export/{project_id}")
async def export_project(
    project_id: str,
    req: ExportRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    导出小说为 TXT/EPUB/DOCX/PDF
    - TXT: 纯文本 + 编码选择
    - EPUB: ebooklib 生成含目录+封面
    - DOCX: python-docx 生成含样式
    - PDF: reportlab 生成含排版
    """
    project = await get_user_project(project_id, user, db)

    chs = await db.execute(
        select(NovelChapter)
        .where(NovelChapter.project_id == project.id)
        .order_by(NovelChapter.chapter_num)
    )
    chapters = chs.scalars().all()

    if not chapters:
        raise HTTPException(400, "项目无章节可导出")

    fmt = req.format.lower()
    if fmt not in EXPORTERS:
        raise HTTPException(400, f"不支持的格式: {fmt}。支持: {', '.join(EXPORTERS.keys())}")

    total_words = sum(ch.word_count for ch in chapters)

    try:
        if fmt == "txt":
            content = EXPORTERS[fmt](chapters, project.title, encoding=req.encoding)
        elif fmt == "epub":
            content = EXPORTERS[fmt](chapters, project.title,
                                     include_cover=req.include_cover,
                                     cover_text=req.cover_text)
        else:
            content = EXPORTERS[fmt](chapters, project.title)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"导出失败: {e}")

    file_size = len(content)
    safe_title = "".join(c for c in project.title if c.isalnum() or c in "._- ").strip() or "novel"
    filename = f"{safe_title}{FILE_EXTENSIONS[fmt]}"

    return StreamingResponse(
        io.BytesIO(content),
        media_type=CONTENT_TYPES[fmt],
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(file_size),
        },
    )


@router.get("/projects/{project_id}/export-info")
async def export_info(
    project_id: str,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """获取导出信息（章节数、字数、支持格式）"""
    project = await get_user_project(project_id, user, db)

    chs = await db.execute(
        select(NovelChapter)
        .where(NovelChapter.project_id == project.id)
        .order_by(NovelChapter.chapter_num)
    )
    chapters = chs.scalars().all()

    total_words = sum(ch.word_count for ch in chapters)

    return {
        "project_id": project_id,
        "title": project.title,
        "total_chapters": len(chapters),
        "total_words": total_words,
        "supported_formats": list(EXPORTERS.keys()),
        "estimated_sizes": {
            "txt": f"{total_words * 3 // 1024} KB (estimated)",
            "epub": f"{total_words * 3 // 2048} KB (estimated)",
            "docx": f"{total_words * 4 // 1024} KB (estimated)",
            "pdf": f"{total_words * 5 // 1024} KB (estimated)",
        },
    }


@router.get("/projects/{project_id}/export")
async def export_project_legacy(project_id: str, format: str = Query(default="txt"),
                                db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    """Legacy TXT export endpoint (backward compatible)."""
    await get_user_project(project_id, user, db)
    chs = await db.execute(select(NovelChapter).where(NovelChapter.project_id == project_id).order_by(NovelChapter.chapter_num))
    chapters = chs.scalars().all()
    if format not in EXPORTERS:
        raise HTTPException(400, f"不支持的格式: {format}。支持: {', '.join(EXPORTERS.keys())}")
    project_title = chapters[0].project.title if chapters else ""
    content = EXPORTERS[format](chapters, project_title)
    return StreamingResponse(
        io.BytesIO(content),
        media_type=CONTENT_TYPES.get(format, "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{project_title or "novel"}{FILE_EXTENSIONS.get(format, ".txt")}"'},
    )


@router.get("/projects/{project_id}/versions")
async def chapter_versions(project_id: str, db: AsyncSession = Depends(get_db),
                           user: User = Depends(get_current_user)):
    await get_user_project(project_id, user, db)
    chs = await db.execute(select(NovelChapter).where(NovelChapter.project_id == project_id).order_by(NovelChapter.chapter_num))
    return [{"id": str(c.id), "chapter_num": c.chapter_num, "title": c.title,
             "version_count": len(c.version_history or []),
             "versions": (c.version_history or [])[-5:]} for c in chs.scalars().all()]


@router.get("/dashboard")
async def ops_dashboard(db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    user_pids = [p[0] for p in (await db.execute(
        select(NovelProject.id).where(NovelProject.user_id == user.id)))]
    ch_count = await db.execute(select(func.count(NovelChapter.id)).where(NovelChapter.project_id.in_(user_pids))) if user_pids else None
    word_total = await db.execute(select(func.sum(NovelChapter.word_count)).where(NovelChapter.project_id.in_(user_pids))) if user_pids else None
    status_dist = await db.execute(select(NovelProject.status, func.count(NovelProject.id))
        .where(NovelProject.user_id == user.id).group_by(NovelProject.status))
    return {"total_projects": len(user_pids),
            "total_chapters": (ch_count.scalar() or 0) if ch_count else 0,
            "total_words": (word_total.scalar() or 0) if word_total else 0,
            "status_distribution": {s: c for s, c in status_dist}}
